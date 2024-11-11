# -*- Product under GNU GPL v3 -*-
# -*- Author: E.Aivayan -*-
import uuid

from docx import Document

from app.conf import BASE_DIR
from app.database.postgre.pg_bugs import get_bugs
from app.database.postgre.testcampaign import get_campaign_content
from app.database.utils.combined_results import get_ticket_with_scenarios
from app.schema.campaign_schema import CampaignFull, Scenario, ScenarioInternal, TicketScenario
from app.schema.error_code import ApplicationError, ApplicationErrorCode
from app.schema.postgres_enums import ScenarioStatusEnum, TestResultStatusEnum
from app.schema.rest_enum import DeliverableTypeEnum
from app.utils.project_alias import provide


async def test_plan_from_campaign(campaign: CampaignFull) -> str:
    document = Document()
    document.add_heading("Test Plan", 0)
    document.add_paragraph(
        f"Campaign for {campaign.project_name} " f"in version {campaign.version}",
        style="Subtitle",
    )
    document.add_page_break()
    document.add_heading("Test scope")
    # Create table of tickets with a default column for acceptance criteria
    table = document.add_table(rows=1, cols=3, style="TableGrid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Reference"
    hdr_cells[1].text = "Summary"
    hdr_cells[2].text = "# Acceptance criteria"
    for ticket in campaign.tickets:
        row_cells = table.add_row().cells
        row_cells[0].text = ticket.reference
        row_cells[1].text = ticket.summary

    document.add_heading("Test environment")
    # Add test environment
    document.add_paragraph("Here add data about test environment")

    document.add_heading("Test scope impediments and non-testable items")
    # Create table of ticket with two column for testability and reason
    table = document.add_table(
        rows=1,
        cols=4,
        style="TableGrid",
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Reference"
    hdr_cells[1].text = "Summary"
    hdr_cells[2].text = "Testability"
    hdr_cells[3].text = "Reason"
    for ticket in campaign.tickets:
        row_cells = table.add_row().cells
        row_cells[0].text = ticket.reference
        row_cells[1].text = ticket.summary

    document.add_heading("Test scope estimation")
    # Add table of ticket with one column for estimation
    table = document.add_table(
        rows=1,
        cols=3,
        style="TableGrid",
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Reference"
    hdr_cells[1].text = "Summary"
    hdr_cells[2].text = "Estimation (md)"
    for ticket in campaign.tickets:
        row_cells = table.add_row().cells
        row_cells[0].text = ticket.reference
        row_cells[1].text = ticket.summary
    # Add total estimation
    document.add_paragraph("The total test execution estimation is <your estimation>md.")
    # Add start and end forecast
    document.add_paragraph(
        "We plan a test execution start at <start date> "
        "and with the current estimation expect an end forecast date "
        "on <end date>."
    )

    document.add_heading("Campaign scenario")
    # Create subsection for each tickets with table of scenario
    for ticket in campaign.tickets:
        document.add_heading(
            f"Scenarios for {ticket.reference}",
            2,
        )
        table = document.add_table(
            rows=1,
            cols=5,
            style="TableGrid",
        )
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Scenario id"
        hdr_cells[1].text = "Scenario name"
        hdr_cells[2].text = "Feature name"
        hdr_cells[3].text = "Epic name"
        hdr_cells[4].text = "Steps"
        for scenario in ticket.scenarios:
            row_cells = table.add_row().cells
            row_cells[0].text = scenario.scenario_id
            row_cells[1].text = scenario.name
            row_cells[2].text = scenario.feature_id
            row_cells[3].text = scenario.epic_id
            row_cells[4].text = scenario.steps

    filename = (
        BASE_DIR / "static" / f"Test_Plan_{provide(campaign.project_name)}_" f"{campaign.version}_{uuid.uuid4()}.docx"
    )  # pragma:noqa

    document.save(filename)

    return filename.name


def _compute_status(scenarios: list[Scenario | ScenarioInternal] | None) -> TestResultStatusEnum:
    status = [scenario.status for scenario in scenarios]
    if ScenarioStatusEnum.waiting_fix in status or ScenarioStatusEnum.waiting_answer in status:
        return TestResultStatusEnum.failed
    if (
        all(ScenarioStatusEnum(stat) in [ScenarioStatusEnum.done, ScenarioStatusEnum.cancelled] for stat in status)
        and status
    ):
        return TestResultStatusEnum.passed
    return TestResultStatusEnum.skipped


async def test_exit_report_from_campaign(campaign: CampaignFull) -> str:
    document = Document()
    document.add_heading("Test Exit Report", 0)
    document.add_paragraph(
        f"Campaign for {campaign.project_name} " f"in version {campaign.version}",
        style="Subtitle",
    )
    document.add_page_break()
    document.add_heading("Test campaign overview")
    # Add go/no go sentence

    document.add_heading("Test scope")
    # Create table of ticket
    table = document.add_table(
        rows=1,
        cols=2,
        style="TableGrid",
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Reference"
    hdr_cells[1].text = "Summary"
    for ticket in campaign.tickets:
        row_cells = table.add_row().cells
        row_cells[0].text = ticket.reference
        row_cells[1].text = ticket.summary

    document.add_heading("Test campaign indicators")
    document.add_heading(
        "Environment",
        2,
    )
    # Add template for test environment
    document.add_paragraph(
        "Operating system: ",
        style="List Bullet",
    )
    document.add_paragraph(
        "Browser (version): ",
        style="List Bullet",
    )
    document.add_paragraph(
        "Application environment: ",
        style="List Bullet",
    )

    document.add_heading(
        "Schedule",
        2,
    )
    # Add test campaign start/end dates or leave it blank
    document.add_paragraph(
        "Start date: ",
        style="List Bullet",
    )
    document.add_paragraph(
        "End date: ",
        style="List Bullet",
    )
    document.add_paragraph(
        "End reason: ",
        style="List Bullet",
    )

    document.add_heading(
        "Impediment",
        2,
    )
    # Add template for impediment section

    document.add_heading(
        "Test result summary",
    )
    # Create table of ticket with computed status based on scenarios status
    table = document.add_table(
        rows=1,
        cols=4,
        style="TableGrid",
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Reference"
    hdr_cells[1].text = "Summary"
    hdr_cells[2].text = "Status"
    hdr_cells[3].text = "Comment"
    for ticket in campaign.tickets:
        row_cells = table.add_row().cells
        row_cells[0].text = ticket.reference
        row_cells[1].text = ticket.summary
        row_cells[2].text = _compute_status(ticket.scenarios)

    document.add_heading("Defect status")
    # Create table of defect within the version
    bugs, _ = await get_bugs(
        project_name=campaign.project_name,
        version=campaign.version,
    )
    table = document.add_table(
        rows=1,
        cols=3,
        style="TableGrid",
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Title"
    hdr_cells[1].text = "Criticality"
    hdr_cells[2].text = "Status"
    for bug in bugs:
        row_cells = table.add_row().cells
        row_cells[0].text = bug["title"]
        row_cells[1].text = bug["criticality"]
        row_cells[2].text = bug.status.value

    filename = (
        BASE_DIR
        / "static"
        / (f"TER_{provide(campaign.project_name)}" f"_{campaign.version}_" f"{campaign.occurrence}{uuid.uuid4()}.docx")
    )  # pragma:noqa

    document.save(filename)

    return filename.name


async def evidence_from_ticket(ticket: TicketScenario) -> str:
    document = Document()
    document.add_heading(
        "Test Evidence",
        0,
    )

    document.add_paragraph(
        f"Ticket {ticket.reference} test execution evidence",
        style="Subtitle",
    )

    document.add_page_break()
    document.add_heading("Scenarios")
    for scenario in ticket.scenarios:
        document.add_heading(scenario.name, 2)
        document.add_paragraph(f"Scenario id is {scenario.scenario_id}.")
        document.add_paragraph(scenario.steps)

    document.add_heading("Test conditions")
    document.add_paragraph(
        "Operating system: ",
        style="List Bullet",
    )
    document.add_paragraph(
        "Browser (version): ",
        style="List Bullet",
    )
    document.add_paragraph(
        "Application environment: ",
        style="List Bullet",
    )
    document.add_paragraph(
        "Start date: ",
        style="List Bullet",
    )
    document.add_paragraph(
        "End date: ",
        style="List Bullet",
    )

    document.add_heading("Prerequisites")

    document.add_heading("Test evidence")
    for scenario in ticket.scenarios:
        document.add_heading(scenario.name, 2)
        document.add_paragraph(f"Scenario id is {scenario.scenario_id}.")
        document.add_paragraph(scenario.steps)

    document.add_heading("Test execution conclusion")

    filename = BASE_DIR / "static" / f"evidence_{ticket.reference}-{uuid.uuid4()}.docx"

    document.save(filename)

    return filename.name


async def campaign_deliverable(
    project_name: str,
    version: str,
    occurrence: str,
    deliverable_type: DeliverableTypeEnum,
    ticket_ref: str = None,
) -> str | ApplicationError:
    # TODO register file
    # file:project_alias:version:occurrence:type
    match deliverable_type:
        case DeliverableTypeEnum.TEST_PLAN:
            campaign = await get_campaign_content(project_name, version, occurrence)
            if isinstance(campaign, ApplicationError):
                return campaign
            filename = await test_plan_from_campaign(campaign)
        case DeliverableTypeEnum.TER:
            campaign = await get_campaign_content(project_name, version, occurrence)
            if isinstance(campaign, ApplicationError):
                return campaign
            filename = await test_exit_report_from_campaign(campaign)
        case DeliverableTypeEnum.EVIDENCE:
            ticket = await get_ticket_with_scenarios(project_name, version, occurrence, ticket_ref)
            if isinstance(ticket, ApplicationError):
                return ticket
            filename = await evidence_from_ticket(ticket)
        case _:
            return ApplicationError(
                error=ApplicationErrorCode.value_error,
                message="This value is not implemented yet.",
            )
    return filename
